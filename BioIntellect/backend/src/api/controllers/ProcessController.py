from .BaseController import BaseController

import os
from .ProjectController import projectController
from langchain_community.document_loaders import TextLoader
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document


class _DocxLoader:
    """Minimal loader for .docx files using python-docx."""

    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> list:
        try:
            import docx  # python-docx
        except ImportError as exc:
            raise ImportError(
                "python-docx is required to load .docx files. "
                "Install it with: pip install python-docx"
            ) from exc

        doc = docx.Document(self.file_path)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        return [Document(page_content=text, metadata={"source": self.file_path})]


class ProcessController(BaseController) :

    def __init__(self,project_id:str):
        super().__init__()
        self.project_id=project_id
        self.project_path=projectController().get_project_path(project_id=project_id)
    def get_file_extension(self,file_id:str):
        return os.path.splitext(file_id)[1].lower()
    def get_file_loader(self,file_id:str):
        file_ext=self.get_file_extension(file_id=file_id)
        base_dir = os.path.realpath(self.project_path)
        file_path=os.path.realpath(os.path.join(base_dir,file_id))

        if not file_path.startswith(base_dir + os.sep):
            raise ValueError("file_id resolves outside the allowed directory")

        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"File does not exist: {file_path}")

        if file_ext in [".txt", ".md", ".csv", ".json"]:
            return TextLoader(file_path, encoding="utf-8", autodetect_encoding=True)
        elif file_ext == ".pdf":
            return PyPDFLoader(file_path)
        elif file_ext == ".docx":
            return _DocxLoader(file_path)
        else:
            raise ValueError(f"unsupported file type: {file_ext}")
        


    def get_file_content(self,file_id:str):
        loader=self.get_file_loader(file_id=file_id)
        return loader.load()

    @staticmethod
    def adaptive_chunk_params(total_chars: int) -> tuple[int, int]:
        """Pick chunk_size/overlap based on document length.

        Rationale: tiny files shouldn't be split into one micro-chunk that
        loses retrieval granularity, and huge files shouldn't be split into
        thousands of 500-char chunks that bloat the index and embedding cost.
        Overlap is kept at ~15% so neighboring chunks share boundary context.
        """
        if total_chars <= 0:
            return 500, 75
        if total_chars < 4_000:          # ~1-2 pages
            chunk_size = 400
        elif total_chars < 20_000:       # ~5-10 pages
            chunk_size = 700
        elif total_chars < 80_000:       # ~20-40 pages
            chunk_size = 1000
        elif total_chars < 300_000:      # ~80-150 pages
            chunk_size = 1400
        else:                            # very large docs / books
            chunk_size = 1800
        overlap_size = max(50, int(chunk_size * 0.15))
        return chunk_size, overlap_size

    def process_file_content(self, file_content: list, file_id: str,
                             chunk_size: int | None = None,
                             overlap_size: int | None = None):
        file_content_text = [doc.page_content for doc in file_content]
        file_content_metadata = [dict(doc.metadata or {}) for doc in file_content]

        # Adaptive sizing when the caller did not pin explicit values.
        if chunk_size is None or overlap_size is None:
            total_chars = sum(len(t or "") for t in file_content_text)
            auto_cs, auto_os = self.adaptive_chunk_params(total_chars)
            chunk_size = chunk_size or auto_cs
            overlap_size = overlap_size if overlap_size is not None else auto_os
        if overlap_size >= chunk_size:
            overlap_size = max(0, chunk_size // 10)

        # Propagate filename/source into every chunk's metadata so the LLM
        # (and the UI) can cite which document each retrieved chunk came from.
        display_name = os.path.basename(file_id or "") or file_id
        for meta in file_content_metadata:
            meta.setdefault("source_file_id", file_id)
            meta.setdefault("file_name", display_name)

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=overlap_size,
            length_function=len,
        )
        chunks = text_splitter.create_documents(
            file_content_text,
            metadatas=file_content_metadata,
        )
        for idx, chunk in enumerate(chunks):
            if chunk.metadata is None:
                chunk.metadata = {}
            chunk.metadata.setdefault("source_file_id", file_id)
            chunk.metadata.setdefault("file_name", display_name)
            chunk.metadata["chunk_index"] = idx
        return chunks